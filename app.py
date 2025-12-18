import tempfile
import json
import os
import time
from datetime import datetime
from functools import lru_cache
from flask import Flask, render_template, request, send_file, redirect, url_for, session, jsonify
import yfinance as yf
import finnhub
from docx import Document
import io
import threading
from requests.exceptions import HTTPError, RequestException

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your_secret_key_change_in_production')

# Finnhub client setup
finnhub_client = finnhub.Client(api_key=os.environ.get('FINNHUB_API_KEY', 'cs97jkhr01qoa9gbio60cs97jkhr01qoa9gbio6g'))

# In-memory storage for analysis progress
analysis_status = {}
analysis_results = {}

def fetch_with_retry(func, max_retries=5, initial_delay=3):
    """Retry wrapper with exponential backoff"""
    for attempt in range(max_retries):
        try:
            result = func()
            if result:  # If we got data, return immediately
                return result
            # If no data but no error, retry
            if attempt < max_retries - 1:
                delay = initial_delay * (1.5 ** attempt)
                print(f"No data returned, waiting {delay:.1f}s before retry {attempt + 1}/{max_retries}")
                time.sleep(delay)
        except HTTPError as e:
            if e.response.status_code == 429:  # Too Many Requests
                if attempt < max_retries - 1:
                    delay = initial_delay * (2 ** attempt)
                    print(f"Rate limited (429), waiting {delay}s before retry {attempt + 1}/{max_retries}")
                    time.sleep(delay)
                else:
                    print(f"Max retries reached after 429 error")
                    return None
            else:
                print(f"HTTP Error {e.response.status_code}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(initial_delay)
                else:
                    return None
        except RequestException as e:
            print(f"Request exception: {e}")
            if attempt < max_retries - 1:
                time.sleep(initial_delay)
            else:
                return None
        except Exception as e:
            print(f"Unexpected error: {e}")
            if attempt < max_retries - 1:
                time.sleep(initial_delay)
            else:
                return None
    return None

def get_stock_data(ticker):
    """Fetch stock data with retry logic - no caching to avoid stale data"""
    def fetch():
        try:
            print(f"Fetching data for {ticker}...")
            stock = yf.Ticker(ticker)
            
            # Force fresh data fetch
            info = stock.info
            
            # Validate we got meaningful data
            if info and len(info) > 5:  # Should have more than just basic fields
                print(f"Successfully fetched {len(info)} fields for {ticker}")
                return info
            else:
                print(f"Insufficient data for {ticker}, got {len(info) if info else 0} fields")
                return None
        except Exception as e:
            print(f"Error in get_stock_data for {ticker}: {e}")
            return None
    
    result = fetch_with_retry(fetch, max_retries=5, initial_delay=3)
    return result if result else {}

def get_forward_PE(ticker):
    try:
        info = get_stock_data(ticker)
        if not info:
            print(f"No info available for forward PE of {ticker}")
            return None
        forward_pe = info.get('forwardPE', None)
        if forward_pe:
            print(f"{ticker} forward PE: {forward_pe}")
        else:
            print(f"{ticker} forward PE not available in data")
        return forward_pe
    except Exception as e:
        print(f"Error fetching forward PE for {ticker}: {e}")
        return None

def get_price_target_data(ticker):
    try:
        info = get_stock_data(ticker)
        if not info:
            print(f"No info available for price targets of {ticker}")
            return None, None, None
        mean_price = info.get('targetMeanPrice', None)
        high_price = info.get('targetHighPrice', None)
        low_price = info.get('targetLowPrice', None)
        print(f"{ticker} price targets - Mean: {mean_price}, High: {high_price}, Low: {low_price}")
        return mean_price, high_price, low_price
    except Exception as e:
        print(f"Error fetching price targets for {ticker}: {e}")
        return None, None, None

def get_analyst_ratings_finnhub(ticker):
    def fetch():
        return finnhub_client.recommendation_trends(ticker)
    
    try:
        recommendations = fetch_with_retry(fetch, max_retries=3, initial_delay=2)
        if not recommendations:
            print(f"No analyst ratings for {ticker}")
            return None, 0
        
        latest_recommendation = recommendations[0]
        total_ratings = sum([
            latest_recommendation.get('buy', 0),
            latest_recommendation.get('hold', 0),
            latest_recommendation.get('sell', 0),
            latest_recommendation.get('strongBuy', 0),
            latest_recommendation.get('strongSell', 0)
        ])
        
        buy_ratings = latest_recommendation.get('buy', 0) + latest_recommendation.get('strongBuy', 0)
        buy_percentage = (buy_ratings / total_ratings) * 100 if total_ratings > 0 else 0
        print(f"{ticker} analyst ratings: {buy_percentage:.1f}% buy ({buy_ratings}/{total_ratings})")
        return buy_percentage, total_ratings
    except Exception as e:
        print(f"Error fetching analyst ratings for {ticker}: {e}")
        return None, 0

def get_ownership_data(ticker):
    try:
        info = get_stock_data(ticker)
        if not info:
            print(f"No info available for ownership of {ticker}")
            return None, None, None
            
        insiders = info.get('heldPercentInsiders', None)
        institutions = info.get('heldPercentInstitutions', None)
        
        print(f"{ticker} ownership - Insiders: {insiders}, Institutions: {institutions}")
        
        if insiders is None or institutions is None:
            insider_ownership = insiders * 100 if insiders is not None else None
            institutional_ownership = institutions * 100 if institutions is not None else None
            return None, insider_ownership, institutional_ownership
        
        insider_ownership = insiders * 100
        institutional_ownership = institutions * 100
        total_ownership = insider_ownership + institutional_ownership
        
        return total_ownership, insider_ownership, institutional_ownership
    except Exception as e:
        print(f"Error fetching ownership data for {ticker}: {e}")
        return None, None, None

def check_all_conditions(ticker, benchmark_forward_PE):
    print(f"\n{'='*60}")
    print(f"Starting analysis for {ticker}")
    print(f"{'='*60}")
    
    info = get_stock_data(ticker)
    ticker_results = [ticker]
    details = [ticker]
    conditions_met = True
    
    # 1. Analyst Ratings Check
    print(f"\n[{ticker}] Checking analyst ratings...")
    buy_percentage, buy_count = get_analyst_ratings_finnhub(ticker)
    if buy_percentage is None:
        ticker_results.append('❌ (Ratings Unavailable)')
        details.append(f"Buy Ratings: Data unavailable")
        conditions_met = False
    else:
        details.append(f"Buy Ratings: {buy_count} analysts, Buy Percentage: {buy_percentage:.2f}%")
        if buy_percentage < 70:
            ticker_results.append('❌')
            conditions_met = False
        else:
            ticker_results.append('✔️')
    
    # 2. Market Cap Check
    print(f"[{ticker}] Checking market cap...")
    market_cap = info.get('marketCap', 0) / 1e9 if info.get('marketCap') else 0
    details.append(f"Market Cap: ${market_cap:.2f}B")
    if market_cap < 100:
        ticker_results.append('❌')
        conditions_met = False
    else:
        ticker_results.append('✔️')
    
    # 3. Price Target Check
    print(f"[{ticker}] Checking price targets...")
    mean_price, high_price, low_price = get_price_target_data(ticker)
    current_price = info.get('currentPrice', 0) if info.get('currentPrice') else 0
    details.append(f"Price Targets: Current=${current_price:.2f}, Mean=${mean_price}, High=${high_price}, Low=${low_price}")
    if mean_price is None or mean_price <= current_price * 1.15 or low_price < current_price:
        ticker_results.append('❌')
        conditions_met = False
    else:
        ticker_results.append('✔️')
    
    # 4. Forward PE Check
    print(f"[{ticker}] Checking forward P/E...")
    forward_PE_next_year = get_forward_PE(ticker)
    details.append(f"Forward PE: {forward_PE_next_year if forward_PE_next_year is not None else 'N/A'} (Benchmark: {benchmark_forward_PE})")
    if forward_PE_next_year is None:
        ticker_results.append('❌ (Missing Forward PE)')
        conditions_met = False
    elif forward_PE_next_year >= benchmark_forward_PE:
        ticker_results.append('❌')
        conditions_met = False
    else:
        ticker_results.append('✔️')
    
    # 5. Ownership Check
    print(f"[{ticker}] Checking ownership...")
    total_ownership, insider_ownership, institutional_ownership = get_ownership_data(ticker)
    if total_ownership is None:
        details.append(f"Ownership: Data unavailable (Insider: {insider_ownership if insider_ownership else 'N/A'}%, Institutional: {institutional_ownership if institutional_ownership else 'N/A'}%)")
        ticker_results.append('❌ (Ownership Data Unavailable)')
        conditions_met = False
    else:
        details.append(f"Ownership: Total {total_ownership:.2f}%, Insider: {insider_ownership:.2f}%, Institutional: {institutional_ownership:.2f}%")
        if total_ownership < 70:
            ticker_results.append('❌')
            conditions_met = False
        else:
            ticker_results.append('✔️')
    
    print(f"\n[{ticker}] Analysis complete - Conditions met: {conditions_met}")
    print(f"{'='*60}\n")
    
    return ticker_results, conditions_met, details

def process_tickers_background(job_id, tickers, benchmark_forward_PE):
    """Background worker function"""
    total = len(tickers)
    buy_tickers = []
    results = []
    all_details = []
    
    analysis_status[job_id] = {
        'status': 'processing',
        'progress': 0,
        'total': total,
        'current': 0,
        'start_time': time.time()
    }
    
    print(f"\n{'#'*80}")
    print(f"# Starting Job {job_id}: {total} tickers")
    print(f"{'#'*80}\n")
    
    for i, ticker in enumerate(tickers):
        try:
            print(f"\n[Job {job_id}] Processing {i+1}/{total}: {ticker}")
            ticker_results, conditions_met, details = check_all_conditions(ticker, benchmark_forward_PE)
            results.append(ticker_results)
            all_details.append(details)
            if conditions_met:
                buy_tickers.append(ticker)
                print(f"✓ {ticker} meets all conditions!")
            
            # Update progress
            analysis_status[job_id]['current'] = i + 1
            analysis_status[job_id]['progress'] = int(((i + 1) / total) * 100)
            
            # Rate limiting: longer delays to avoid 429 errors
            if i < total - 1:
                delay = 2.5  # 2.5 seconds between each ticker
                print(f"Waiting {delay}s before next ticker...")
                time.sleep(delay)
                
        except Exception as e:
            print(f"[Job {job_id}] Error processing {ticker}: {e}")
            results.append([ticker, '❌', '❌', '❌', '❌', '❌'])
            all_details.append([ticker, f"Error: {str(e)}"])
    
    # Save results
    analysis_results[job_id] = {
        'results': results,
        'buy_tickers': buy_tickers,
        'all_details': all_details,
        'completed_at': datetime.now().isoformat()
    }
    
    analysis_status[job_id]['status'] = 'completed'
    analysis_status[job_id]['progress'] = 100
    
    print(f"\n{'#'*80}")
    print(f"# Job {job_id} Complete!")
    print(f"# Found {len(buy_tickers)} BUY candidates out of {total} stocks")
    print(f"{'#'*80}\n")

def generate_word_report(results, details, buy_tickers):
    doc = Document()
    doc.add_heading('Stock Analysis Report', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Stocks Analyzed: {len(results)}')
    doc.add_paragraph(f'Stocks Meeting All Criteria: {len(buy_tickers)}')
    doc.add_paragraph()
    
    doc.add_heading('Buy Tickers', level=1)
    if buy_tickers:
        doc.add_paragraph(', '.join(buy_tickers))
    else:
        doc.add_paragraph("No tickers met all conditions.")
    doc.add_paragraph()
    
    doc.add_heading('Detailed Analysis', level=1)
    for ticker_details, ticker_results in zip(details, results):
        doc.add_heading(f"{ticker_details[0]}", level=2)
        for detail in ticker_details[1:]:
            doc.add_paragraph(detail)
        doc.add_paragraph(f"Results: {' '.join(ticker_results[1:])}")
        doc.add_paragraph()
    
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    return word_file

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    tickers_input = request.form.get('tickers', '').strip()
    benchmark_forward_PE_input = request.form.get('benchmark_forward_PE', '').strip()
    
    if not tickers_input or not benchmark_forward_PE_input:
        return "Please provide both tickers and benchmark forward PE.", 400
    
    tickers = [t.strip().upper() for t in tickers_input.split(',') if t.strip()]
    
    try:
        benchmark_forward_PE = float(benchmark_forward_PE_input)
    except ValueError:
        return "Please enter a valid numeric benchmark forward PE.", 400
    
    # Generate unique job ID
    job_id = f"job_{int(time.time())}_{len(tickers)}"
    
    # Start background processing
    thread = threading.Thread(
        target=process_tickers_background,
        args=(job_id, tickers, benchmark_forward_PE)
    )
    thread.daemon = True
    thread.start()
    
    # Redirect to progress page
    return redirect(url_for('progress', job_id=job_id))

@app.route('/progress/<job_id>')
def progress(job_id):
    return render_template('progress.html', job_id=job_id)

@app.route('/api/status/<job_id>')
def get_status(job_id):
    """API endpoint for progress updates"""
    if job_id in analysis_status:
        return jsonify(analysis_status[job_id])
    else:
        return jsonify({'status': 'not_found'}), 404

@app.route('/results/<job_id>')
def results(job_id):
    if job_id not in analysis_results:
        return "Analysis not found or still in progress.", 404
    
    data = analysis_results[job_id]
    return render_template('result.html', 
                         buy_tickers=data['buy_tickers'], 
                         results=data['results'], 
                         all_details=data['all_details'],
                         job_id=job_id)

@app.route('/download_word/<job_id>')
def download_word(job_id):
    if job_id not in analysis_results:
        return "Analysis not found.", 404
    
    data = analysis_results[job_id]
    word_file = generate_word_report(
        data['results'], 
        data['all_details'], 
        data['buy_tickers']
    )
    
    return send_file(
        word_file,
        as_attachment=True,
        download_name=f'sp500_analysis_{job_id}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
